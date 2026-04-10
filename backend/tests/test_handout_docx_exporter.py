import base64
import sys
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.services.handout_docx_exporter import HandoutDocxExporter


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z0uoAAAAASUVORK5CYII="
)


def extract_docx_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(part for part in parts if part)


def test_build_reading_grade_docx_embeds_option_images(tmp_path: Path) -> None:
    static_root = tmp_path / "static"
    image_dir = static_root / "images" / "options"
    image_dir.mkdir(parents=True, exist_ok=True)
    image_path = image_dir / "test_option.png"
    image_path.write_bytes(PNG_1X1)

    exporter = HandoutDocxExporter(static_root=static_root)
    handout = {
        "grade": "初二",
        "edition": "teacher",
        "topics": [{"topic": "太空探索", "passage_count": 1, "recent_years": [2026]}],
        "content": [
            {
                "topic": "太空探索",
                "part1_article_sources": [],
                "part2_vocabulary": [],
                "part3_passages": [
                    {
                        "type": "C",
                        "title": "Black Holes",
                        "source": {"year": 2026, "region": "北京", "grade": "初二", "exam_type": "期中"},
                        "content": "A short reading passage.",
                        "questions": [
                            {
                                "number": 39,
                                "text": "Choose the best answer.",
                                "options": {
                                    "A": "Plain text option",
                                    "B": "Option with image\n[IMAGE:/static/images/options/test_option.png]",
                                },
                                "correct_answer": "B",
                                "explanation": "Because it matches the passage.",
                            }
                        ],
                    }
                ],
            }
        ],
    }

    file_bytes = exporter.build_reading_grade_docx(handout)

    with ZipFile(BytesIO(file_bytes)) as archive:
        media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
        assert media_files, "导出的 docx 应包含嵌入图片"

    document = Document(BytesIO(file_bytes))
    full_text = extract_docx_text(document)
    assert "Option with image" in full_text
    assert "[IMAGE:" not in full_text


def test_build_cloze_grade_docx_uses_chinese_point_labels_for_rejection_codes() -> None:
    exporter = HandoutDocxExporter()
    handout = {
        "grade": "初三",
        "edition": "teacher",
        "topics": [{"topic": "人物故事", "passage_count": 1, "recent_years": [2025]}],
        "content": [
            {
                "topic": "人物故事",
                "part1_article_sources": [],
                "part2_vocabulary": [],
                "part3_points_by_type": {},
                "part4_passages": [
                    {
                        "source": {"year": 2025, "region": "北京", "grade": "初三", "exam_type": "中考"},
                        "content": "He (1) hard every day.",
                        "points": [
                            {
                                "blank_number": 1,
                                "correct_answer": "A",
                                "correct_word": "worked",
                                "options": {
                                    "A": "worked",
                                    "B": "played",
                                    "C": "slept",
                                    "D": "waited",
                                },
                                "rejection_points": [
                                    {
                                        "option_word": "played",
                                        "rejection_code": "A1",
                                    }
                                ],
                                "explanation": "第一层（A1）看上下文信息，因此C2固定搭配更合理。",
                                "tips": "先用A1定位语境，再结合C2判断搭配。",
                            }
                        ],
                    }
                ],
            }
        ],
    }

    file_bytes = exporter.build_cloze_grade_docx(handout)
    document = Document(BytesIO(file_bytes))
    full_text = extract_docx_text(document)
    assert "上下文语义推断" in full_text
    assert "固定搭配更合理" in full_text
    assert "A1" not in full_text
    assert "C2" not in full_text


def test_build_writing_grade_docx_includes_full_template_content() -> None:
    exporter = HandoutDocxExporter()
    handout = {
        "grade": "初二",
        "edition": "teacher",
        "groups": [
            {
                "group_category": {"name": "应用文"},
                "sections": [
                    {
                        "summary": {
                            "group_name": "应用文",
                            "major_category_name": "介绍说明类",
                            "category_name": "活动介绍",
                            "task_count": 3,
                            "sample_count": 3,
                            "recent_years": [2024, 2025],
                            "applicable_ranges": ["适用于校园活动介绍题"],
                        },
                        "frameworks": [
                            {
                                "title": "三段式框架",
                                "sections": [
                                    {
                                        "name": "开头",
                                        "description": "点明活动背景",
                                        "examples": ["I am writing to introduce our school activity."],
                                    }
                                ],
                            }
                        ],
                        "template_content": (
                            "Dear friends,\n"
                            "I am glad to introduce [activity name] to you.\n\n"
                            "It will be held on [date] at [place].\n"
                            "I hope you can join us."
                        ),
                        "expressions": [],
                        "samples": [],
                    }
                ],
            }
        ],
    }

    file_bytes = exporter.build_writing_grade_docx(handout)
    document = Document(BytesIO(file_bytes))
    full_text = extract_docx_text(document)
    assert "Part 3 模板原文" in full_text
    assert "I am glad to introduce [activity name] to you." in full_text
    assert "It will be held on [date] at [place]." in full_text
