"""
文档解析服务

实现三层解析策略：
1. 规则解析 (80%成功率)
2. LLM辅助 (15%成功率)
3. 人工标注 (5%)
"""
import re
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from docx import Document


class ParseStrategy(Enum):
    """解析策略"""
    RULE = "rule"      # 规则解析
    LLM = "llm"        # LLM辅助
    MANUAL = "manual"  # 人工标注


@dataclass
class ParseResult:
    """解析结果"""
    success: bool
    data: Optional[Dict] = None
    confidence: float = 0.0
    strategy: ParseStrategy = ParseStrategy.RULE
    error: Optional[str] = None


@dataclass
class PassageExtract:
    """提取的文章"""
    marker: str  # A, B, C, D
    content: str
    questions: List[Dict] = field(default_factory=list)


class DocxParser:
    """文档解析器"""

    # 阅读理解关键词
    READING_KEYWORDS = [
        "阅读理解", "Reading Comprehension", "阅读下列短文",
        "三、阅读", "四、阅读", "Part III Reading", "Part 3 Reading",
        "第三部分 阅读", "第四部分 阅读"
    ]

    # 完形填空关键词
    CLOZE_KEYWORDS = [
        "完形填空", "Cloze", "完型填空", "完形",
        "二、完形", "Part II Cloze", "第二部分 完形"
    ]

    # 作文关键词
    WRITING_KEYWORDS = [
        "作文", "书面表达", "Writing", "写作",
        "书面表达（共", "九、作文", "十、作文"
    ]

    # 文件名解析正则 - 支持全角和半角括号，学期可选
    FILENAME_PATTERN = re.compile(
        r'(\d{4})北京(.+?)(初[一二三])(?:[（(](上|下)[）)])?(期中|期末|一模|二模)英语[（(]?(教师版|学生版|原卷版|解析版)?[）)]?'
    )

    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        self.doc: Optional[Document] = None
        self.full_text: str = ""
        self.paragraphs: List[str] = []

    def load(self) -> bool:
        """加载文档"""
        try:
            self.doc = Document(str(self.file_path))
            self.paragraphs = [p.text for p in self.doc.paragraphs]
            self.full_text = "\n".join(self.paragraphs)
            return True
        except Exception as e:
            return False

    # 北京区县列表（用于区分学校名和区县名）
    BEIJING_REGIONS = [
        "东城", "西城", "海淀", "朝阳", "丰台", "石景山",
        "通州", "顺义", "昌平", "大兴", "房山", "门头沟",
        "怀柔", "平谷", "密云", "延庆"
    ]

    def parse_filename(self) -> Dict:
        """从文件名提取元数据"""
        filename = self.file_path.name

        match = self.FILENAME_PATTERN.search(filename)
        if match:
            # 提取中间部分（可能是区县名或学校名）
            region_or_school = match.group(2).strip()

            # 判断是区县还是学校
            region = None
            school = None

            # 检查是否包含已知区县
            for r in self.BEIJING_REGIONS:
                if r in region_or_school:
                    region = r
                    # 如果字符串比区县名长，剩余部分可能是学校名
                    remaining = region_or_school.replace(r, "").strip()
                    if remaining:
                        school = remaining
                    break

            # 如果不包含任何已知区县，则认为是学校名
            if region is None:
                school = region_or_school

            result = {
                "year": int(match.group(1)),
                "region": region,
                "school": school,
                "grade": match.group(3),
                "semester": match.group(4),  # 文件名没有就为 None，不推断
                "exam_type": match.group(5),
                "version": match.group(6) or "学生版"
            }
            return result

        # 尝试从目录结构推断
        return self._parse_from_directory()

    def _infer_semester(self, exam_type: str) -> Optional[str]:
        """根据考试类型推断学期"""
        if not exam_type:
            return None
        # 一模、二模通常在下学期（春季）
        if exam_type in ["一模", "二模"]:
            return "下"
        # 开学考、月考通常在上学期（秋季）
        if exam_type in ["开学考", "月考"]:
            return "上"
        # 期中、期末无法确定，需要更多信息
        return None

    def _parse_from_directory(self) -> Dict:
        """从目录结构推断元数据"""
        parts = self.file_path.parts
        result = {}

        # 遍历路径寻找年份
        for part in parts:
            year_match = re.search(r'(\d{4})', part)
            if year_match:
                result["year"] = int(year_match.group(1))
                break

        # 寻找年级和学期类型
        for part in parts:
            grade_match = re.search(r'(初[一二三])', part)
            if grade_match:
                result["grade"] = grade_match.group(1)

            if "秋季" in part:
                result["semester"] = "上"
                if "期中" in part:
                    result["exam_type"] = "期中"
                elif "期末" in part:
                    result["exam_type"] = "期末"
            elif "春季" in part:
                result["semester"] = "下"
                if "期中" in part:
                    result["exam_type"] = "期中"
                elif "期末" in part:
                    result["exam_type"] = "期末"

        # 寻找区县
        for part in parts:
            # 常见区县
            regions = ["东城", "西城", "海淀", "朝阳", "丰台", "石景山",
                      "通州", "顺义", "昌平", "大兴", "房山", "门头沟",
                      "怀柔", "平谷", "密云", "延庆"]
            for region in regions:
                if region in part:
                    result["region"] = region
                    break

        return result

    def extract_reading_passages(self) -> ParseResult:
        """提取阅读C/D篇 - 三层策略"""
        if not self.doc:
            if not self.load():
                return ParseResult(
                    success=False,
                    confidence=0.0,
                    error="无法加载文档"
                )

        # Layer 1: 规则解析
        result = self._extract_by_rules()

        if result.confidence >= 0.8:
            result.strategy = ParseStrategy.RULE
            return result

        # Layer 2: LLM辅助 (在调用层处理)
        result.strategy = ParseStrategy.LLM
        result.success = False
        result.error = "需要LLM辅助提取"

        return result

    def _extract_by_rules(self) -> ParseResult:
        """Layer 1: 规则解析"""
        try:
            # 查找阅读理解部分
            reading_start = self._find_section_start(self.READING_KEYWORDS)
            if reading_start == -1:
                return ParseResult(
                    success=False,
                    confidence=0.0,
                    error="未找到阅读理解部分"
                )

            # 提取所有文章
            passages = self._extract_all_passages(reading_start)

            if len(passages) >= 2:
                # 取最后两篇（C篇和D篇）
                c_passage = passages[-2]
                d_passage = passages[-1]

                confidence = self._calculate_confidence(passages)

                return ParseResult(
                    success=True,
                    data={
                        "c_passage": c_passage,
                        "d_passage": d_passage,
                        "all_passages": passages
                    },
                    confidence=confidence
                )

            return ParseResult(
                success=False,
                confidence=0.3,
                error=f"只找到{len(passages)}篇文章"
            )

        except Exception as e:
            return ParseResult(
                success=False,
                confidence=0.0,
                error=str(e)
            )

    def _find_section_start(self, keywords: List[str]) -> int:
        """查找指定部分的起始位置"""
        for i, text in enumerate(self.paragraphs):
            text = text.strip()
            for keyword in keywords:
                if keyword in text:
                    return i
        return -1

    def _extract_all_passages(self, start_idx: int) -> List[PassageExtract]:
        """提取所有阅读文章"""
        passages = []
        current_content = []
        current_marker = None
        in_passage = False

        # 文章标记模式
        marker_patterns = [
            re.compile(r'^([A-D])$', re.I),  # 单独的字母A/B/C/D
            re.compile(r'^([A-D])[\.）\s]', re.I),
            re.compile(r'^Passage\s+([A-D])', re.I),
            re.compile(r'^第([一二三四ABCD])篇', re.I),
        ]

        for i, text in enumerate(self.paragraphs[start_idx:], start=start_idx):
            text = text.strip()

            # 检测新文章开始
            new_marker = None
            for pattern in marker_patterns:
                match = pattern.match(text)
                if match:
                    new_marker = match.group(1).upper()
                    if new_marker in ['一', '二', '三', '四']:
                        mapping = {'一': 'A', '二': 'B', '三': 'C', '四': 'D'}
                        new_marker = mapping[new_marker]
                    break

            if new_marker:
                # 保存上一篇文章（仅当内容足够长时才保存）
                if current_content and current_marker:
                    content = "\n".join(current_content)
                    if len(content) >= 200:  # 最小文章长度
                        passages.append(PassageExtract(
                            marker=current_marker,
                            content=content
                        ))
                current_content = []
                current_marker = new_marker
                in_passage = True
                continue

            # 检测题目开始（数字开头）
            if re.match(r'^\d+[\.\、]', text) and in_passage:
                if current_content and current_marker:
                    content = "\n".join(current_content)
                    if len(content) >= 200:
                        passages.append(PassageExtract(
                            marker=current_marker,
                            content=content
                        ))
                current_content = []
                current_marker = None
                in_passage = False
                # 不break，继续寻找下一篇
                continue

            if in_passage and text:
                current_content.append(text)

        # 添加最后一篇
        if current_content and current_marker:
            content = "\n".join(current_content)
            if len(content) >= 200:
                passages.append(PassageExtract(
                    marker=current_marker,
                    content=content
                ))

        return passages

    def _calculate_confidence(self, passages: List[PassageExtract]) -> float:
        """计算解析置信度"""
        if not passages:
            return 0.0

        # 基于文章数量和内容长度计算置信度
        avg_length = sum(len(p.content) for p in passages) / len(passages)

        if avg_length > 300:
            return 0.9
        elif avg_length > 200:
            return 0.7
        elif avg_length > 100:
            return 0.5
        else:
            return 0.3

    def extract_cloze(self) -> ParseResult:
        """提取完形填空"""
        if not self.doc:
            if not self.load():
                return ParseResult(
                    success=False,
                    confidence=0.0,
                    error="无法加载文档"
                )

        # 查找完形填空部分
        cloze_start = self._find_section_start(self.CLOZE_KEYWORDS)
        if cloze_start == -1:
            return ParseResult(
                success=False,
                confidence=0.0,
                error="未找到完形填空部分"
            )

        # 提取完形填空内容
        content = self._extract_section_content(cloze_start)

        if content:
            return ParseResult(
                success=True,
                data={"content": content},
                confidence=0.8
            )

        return ParseResult(
            success=False,
            confidence=0.0,
            error="无法提取完形填空内容"
        )

    def extract_writing(self) -> ParseResult:
        """提取作文题目"""
        if not self.doc:
            if not self.load():
                return ParseResult(
                    success=False,
                    confidence=0.0,
                    error="无法加载文档"
                )

        # 查找作文部分
        writing_start = self._find_section_start(self.WRITING_KEYWORDS)
        if writing_start == -1:
            return ParseResult(
                success=False,
                confidence=0.0,
                error="未找到作文部分"
            )

        # 提取作文内容
        content = self._extract_section_content(writing_start)

        if content:
            return ParseResult(
                success=True,
                data={"content": content},
                confidence=0.8
            )

        return ParseResult(
            success=False,
            confidence=0.0,
            error="无法提取作文内容"
        )

    def _extract_section_content(self, start_idx: int) -> str:
        """提取指定部分的内容"""
        content = []

        for text in self.paragraphs[start_idx + 1:]:
            text = text.strip()

            # 检测是否到达下一个大题
            if re.match(r'^[一二三四五六七八九十]+[、\.]', text):
                break
            if re.match(r'^Part\s+[IVX]+', text, re.I):
                break

            if text:
                content.append(text)

        return "\n".join(content)
