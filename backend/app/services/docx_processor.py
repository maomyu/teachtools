"""
Word 文档处理服务

删除"参考答案"之后的所有内容，保留图片和题目
"""
from docx import Document
from typing import Dict, Optional
from pathlib import Path


class DocxProcessor:
    """Word 文档处理器 - 删除答案部分，保留题目和图片"""

    # 关键词列表（多种可能的写法）
    ANSWER_KEYWORDS = [
        "参考答案", "答案与解析", "参考答案及解析",
        "答案解析", "试题答案", "参考解答",
        "Answer Key", "Answers", "KEY", "Solutions"
    ]

    def remove_answers_section(
        self,
        doc_path: str,
        output_path: str
    ) -> Dict:
        """
        删除"参考答案"之后的所有内容

        Args:
            doc_path: 原文档路径
            output_path: 输出路径

        Returns:
            {
                "success": bool,
                "removed_count": int,
                "answer_section_start": int or None,
                "keyword_found": str or None
            }
        """
        print(f"[DocxProcessor] 开始处理文档: {doc_path}")
        doc = Document(doc_path)
        total_paragraphs = len(doc.paragraphs)
        print(f"[DocxProcessor] 文档共有 {total_paragraphs} 个段落")

        # 找到"参考答案"所在的位置
        answer_start_index = None
        keyword_found = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # 跳过空段落
            if not text:
                continue

            # 检查是否包含关键词
            for keyword in self.ANSWER_KEYWORDS:
                if keyword in text:
                    answer_start_index = i
                    keyword_found = keyword
                    print(f"[DocxProcessor] 在段落 {i} 找到关键词 '{keyword}': {text[:50]}...")
                    break

            if answer_start_index is not None:
                break

        if answer_start_index is None:
            print("[DocxProcessor] 未找到参考答案标记，保留全部内容")
            doc.save(output_path)
            return {
                "success": True,
                "removed_count": 0,
                "answer_section_start": None,
                "keyword_found": None
            }

        # 从该位置开始，删除所有后续段落
        # 注意：需要从后往前删除，否则索引会变化
        paragraphs_to_remove = []
        for i in range(answer_start_index, total_paragraphs):
            paragraphs_to_remove.append(doc.paragraphs[i])

        removed_count = 0
        for para in paragraphs_to_remove:
            try:
                para._element.getparent().remove(para._element)
                removed_count += 1
            except Exception as e:
                print(f"[DocxProcessor] 删除段落失败: {e}")

        # 保存文档
        doc.save(output_path)
        print(f"[DocxProcessor] 成功删除 {removed_count} 个段落，保存到 {output_path}")

        return {
            "success": True,
            "removed_count": removed_count,
            "answer_section_start": answer_start_index,
            "keyword_found": keyword_found
        }

    def get_paragraph_count(self, doc_path: str) -> int:
        """获取文档段落数量"""
        doc = Document(doc_path)
        return len(doc.paragraphs)

    def get_document_text(self, doc_path: str) -> str:
        """获取文档全文"""
        doc = Document(doc_path)
        return '\n'.join([p.text for p in doc.paragraphs])

    def preview_answer_section(self, doc_path: str) -> Dict:
        """
        预览文档中的答案部分位置（不实际删除）

        Returns:
            {
                "found": bool,
                "keyword": str or None,
                "paragraph_index": int or None,
                "preview_text": str or None
            }
        """
        doc = Document(doc_path)

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            for keyword in self.ANSWER_KEYWORDS:
                if keyword in text:
                    return {
                        "found": True,
                        "keyword": keyword,
                        "paragraph_index": i,
                        "preview_text": text[:100]
                    }

        return {
            "found": False,
            "keyword": None,
            "paragraph_index": None,
            "preview_text": None
        }
