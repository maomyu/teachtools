"""
基于LLM的文档解析服务

使用通义千问qwen-long模型进行文档理解
"""
import os
import json
import re
from difflib import SequenceMatcher
from typing import Dict, List, Optional
from dataclasses import dataclass, field
from docx import Document

from app.config import settings
from app.services.dashscope_runtime import async_chat_completion, async_upload_file


@dataclass
class LLMParseResult:
    """LLM解析结果"""
    success: bool
    passages: List[Dict] = None  # [{"type": "C", "content": "...", "word_count": 300}, ...]
    cloze: Dict = None  # {"found": True, "content_with_blanks": "...", "blanks": [...]}
    metadata: Dict = None  # {"year": 2022, "region": "东城", ...}
    error: str = None
    raw_response: str = None


@dataclass
class WritingExtractTask:
    """单道作文小题提取结果。"""
    content: str = ""                  # 作文题目完整内容
    requirements: str = ""             # 具体要求
    word_limit: str = ""               # 字数限制
    writing_type: str = "其他"         # 应用文/记叙文/其他
    application_type: str = ""         # 应用文子类型
    task_label: str = ""               # 题目① / 题目② / 第一题 ...


@dataclass
class WritingExtractResult:
    """作文提取结果（支持一道作文大题下的多道小题）。"""
    success: bool
    tasks: List[WritingExtractTask] = field(default_factory=list)
    error: str = ""
    raw_response: str = ""

    @property
    def content(self) -> str:
        return self.tasks[0].content if self.tasks else ""

    @property
    def requirements(self) -> str:
        return self.tasks[0].requirements if self.tasks else ""

    @property
    def word_limit(self) -> str:
        return self.tasks[0].word_limit if self.tasks else ""

    @property
    def writing_type(self) -> str:
        return self.tasks[0].writing_type if self.tasks else "其他"

    @property
    def application_type(self) -> str:
        return self.tasks[0].application_type if self.tasks else ""


def _normalize_writing_task_label(label: str) -> str:
    """将题目①/第一题/Task 1 等标签标准化，便于去重比对。"""
    normalized = re.sub(r"[\s：:]", "", label or "").strip().lower()
    replacements = {
        "题目1": "题目①",
        "题目一": "题目①",
        "第一题": "题目①",
        "task1": "题目①",
        "作文1": "题目①",
        "作文一": "题目①",
        "题目2": "题目②",
        "题目二": "题目②",
        "第二题": "题目②",
        "task2": "题目②",
        "作文2": "题目②",
        "作文二": "题目②",
    }
    return replacements.get(normalized, normalized)


def _normalize_writing_task_text(text: str, task_label: str = "") -> str:
    """抽取作文小题的核心题意文本，弱化共享说明和格式差异。"""
    normalized = (text or "").strip()
    if not normalized:
        return ""

    normalized = normalized.replace("（", "(").replace("）", ")").replace("“", '"').replace("”", '"')
    normalized = normalized.replace("‘", "'").replace("’", "'").replace("　", " ")

    label = _normalize_writing_task_label(task_label)
    if label:
        normalized = re.sub(
            rf"(?is)^.*?(?:{re.escape(label)}|{re.escape(task_label.strip())})[\s：:]*",
            "",
            normalized,
            count=1,
        )

    lead_cues = (
        "假设你是", "假如你是", "假定你是", "某英文网站", "你校英语社团", "你们学校",
        "请用英语", "请你用英语", "请根据", "从下面两个题目中任选一题",
    )
    earliest_index = -1
    for cue in lead_cues:
        index = normalized.find(cue)
        if index == -1:
            continue
        if earliest_index == -1 or index < earliest_index:
            earliest_index = index
    if earliest_index > 0:
        normalized = normalized[earliest_index:]

    normalized = re.sub(r"(?is)\|.*?\|", " ", normalized)
    normalized = re.sub(
        r"(?is)(提示词语|提示问题|写作提示|参考答案|范文|优秀范文|参考范文|参考例文|评分标准).*",
        "",
        normalized,
    )
    normalized = re.sub(
        r"(?is)书面表达.*?请不要写出(?:真实的)?校名和姓名[。.]?",
        "",
        normalized,
    )
    normalized = re.sub(r"(?is)文中已给出内容不计入总词数[。.]?", "", normalized)
    normalized = re.sub(r"(?is)所给提示词语仅供选用[。.]?", "", normalized)
    normalized = re.sub(r"(?is)\(共\d+分\)", "", normalized)
    normalized = re.sub(r"[\s\.,;:!?，。；：！？\"'“”‘’()（）【】\[\]_\-—~|]+", "", normalized)
    return normalized.lower()


def dedupe_writing_tasks(tasks: List[WritingExtractTask]) -> List[WritingExtractTask]:
    """对作文小题做语义级去重，保留一张卷里的真实多题，过滤重复提取。"""
    deduped: List[WritingExtractTask] = []
    normalized_cache: List[str] = []

    for task in tasks:
        content = (task.content or "").strip()
        if not content:
            continue

        normalized_text = _normalize_writing_task_text(content, task.task_label)
        normalized_label = _normalize_writing_task_label(task.task_label)
        is_duplicate = False

        for existing_task, existing_text in zip(deduped, normalized_cache):
            existing_label = _normalize_writing_task_label(existing_task.task_label)
            if normalized_label and existing_label and normalized_label != existing_label:
                continue
            if normalized_text and existing_text:
                if (
                    normalized_text == existing_text
                    or normalized_text in existing_text
                    or existing_text in normalized_text
                ):
                    is_duplicate = True
                    break
                similarity = SequenceMatcher(None, normalized_text[:1200], existing_text[:1200]).ratio()
                if similarity >= 0.93:
                    is_duplicate = True
                    break

        if is_duplicate:
            continue

        deduped.append(task)
        normalized_cache.append(normalized_text)

    return deduped


class LLMDocumentParser:
    """基于LLM的文档解析器"""

    API_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
    UPLOAD_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1/files"

    # 文档理解Prompt
    EXTRACT_PROMPT = """你是一个英语试卷解析专家。请分析这份英语试卷文档，提取以下信息：

## 需要提取的信息

### 1. 试卷元数据
从文件名和文档内容中提取：
- year: 年份（如2022）
- region: 区县（如东城、西城、海淀等）
- grade: 年级（初一/初二/初三）
- semester: 学期（上/下）
- exam_type: 考试类型（期中/期末/一模/二模）
- version: 版本（教师版/学生版）

### 2. 阅读理解C篇和D篇
提取阅读理解部分的C篇和D篇，每篇包含：
- passage_type: "C" 或 "D"
- content: 文章完整正文（不要包含题目）
- word_count: 词数（英文单词数）
- questions: 该文章对应的题目列表

### 3. 题目提取
对于每篇文章，提取对应的阅读理解题目：
- question_number: 题号（如31、32、33等）
- question_text: 题目内容（题干）
- options: 选项对象，格式为 {"A": "选项A内容", "B": "选项B内容", ...}
  - 注意：有些题目只有3个选项（A/B/C），这是正常的
  - 必须提取文档中存在的所有选项，确保至少有一个非空选项
  - 如果某个选项不存在，不要包含在options中或设为空字符串
  - **图片选项**：如果选项内容是图片（无法提取文字），填写 "[IMAGE]" 占位符
  - **绝对不要**把思维导图、流程图、示意图、表格图等图片型选项留空，必须填写 "[IMAGE]"
- has_image_options: 布尔值，表示该题目是否有图片选项（true/false）
- expected_image_count: 整数，预期图片选项的数量（0-4）
- correct_answer: 正确答案（A/B/C/D，如果是教师版能找到答案的话）
- answer_explanation: 答案解析（如果是教师版能找到解析的话）

**重要**：
- 普通文本选项：options 必须至少包含一个有效的选项内容，不能全部为空
- 图片选项：如果选项是图片，填写 "[IMAGE]" 占位符，并设置 has_image_options=true
- 混合选项：部分文本、部分图片的情况也要正确标记
- 如果无法识别图片内容，也不要输出空字符串，统一输出 "[IMAGE]"
- 如果是教师版，且原题已经给出了答案或解析，必须优先提取原题答案/原题解析，不要改写成你自己的版本
- 如果是开放题且教师版给出了示例答案，也要优先保留原题示例答案
- 只有在教师版明确没有答案/解析时，才允许补充简短的“参考答案”或“参考解析”

### 4. 完形填空
提取完形填空部分（如果存在）：
- found: 是否找到完形填空（true/false）
- content_with_blanks: 带空格标记的原文（保留空格编号如①②③或(1)(2)(3)）
- content_full: 填入正确答案后的完整文章（如果教师版有答案）
- word_count: 词数（完整文章的英文单词数）
- blanks: 空格列表，每个空格包含：
  - blank_number: 空格编号（1-12）
  - options: 四个选项 {"A": "...", "B": "...", "C": "...", "D": "..."}
  - correct_answer: 正确答案（A/B/C/D，如果教师版有）
  - correct_word: 正确答案对应的词

## 注意事项
- 如果是教师版试卷，注意区分题目、学生版内容和答案解析
- 教师版中的原题答案、示例答案、解答内容优先级最高，优先保留原文信息
- 题目通常紧跟在文章后面，题号一般是连续的
- 确保选项内容完整，不要截断
- 如果找不到答案或解析，对应字段可以省略
- 完形填空通常在阅读理解之前，一般有12个空格

## 输出格式

请严格按照以下JSON格式输出，不要添加任何其他文字：

```json
{
    "metadata": {
        "year": 2022,
        "region": "东城",
        "grade": "初二",
        "semester": "下",
        "exam_type": "期末",
        "version": "教师版"
    },
    "passages": [
        {
            "passage_type": "C",
            "content": "完整文章正文内容...",
            "word_count": 350,
            "questions": [
                {
                    "question_number": 31,
                    "question_text": "题目内容...",
                    "options": {
                        "A": "选项A内容",
                        "B": "选项B内容",
                        "C": "选项C内容",
                        "D": "选项D内容"
                    },
                    "has_image_options": false,
                    "expected_image_count": 0,
                    "correct_answer": "A",
                    "answer_explanation": "答案解析内容..."
                },
                {
                    "question_number": 32,
                    "question_text": "这道题的选项是图片...",
                    "options": {
                        "A": "[IMAGE]",
                        "B": "[IMAGE]",
                        "C": "[IMAGE]",
                        "D": "[IMAGE]"
                    },
                    "has_image_options": true,
                    "expected_image_count": 4,
                    "correct_answer": "B",
                    "answer_explanation": "答案解析内容..."
                }
            ]
        },
        {
            "passage_type": "D",
            "content": "完整文章正文内容...",
            "word_count": 400,
            "questions": [
                {
                    "question_number": 36,
                    "question_text": "题目内容...",
                    "options": {
                        "A": "选项A内容",
                        "B": "选项B内容",
                        "C": "选项C内容",
                        "D": "选项D内容"
                    },
                    "has_image_options": false,
                    "expected_image_count": 0,
                    "correct_answer": "B",
                    "answer_explanation": "答案解析内容..."
                }
            ]
        }
    ],
    "cloze": {
        "found": true,
        "content_with_blanks": "I have always ① reading. I remember the day when I ② my first book...",
        "content_full": "I have always enjoyed reading. I remember the day when I bought my first book...",
        "word_count": 280,
        "blanks": [
            {
                "blank_number": 1,
                "options": {"A": "hated", "B": "disliked", "C": "enjoyed", "D": "avoided"},
                "correct_answer": "C",
                "correct_word": "enjoyed"
            },
            {
                "blank_number": 2,
                "options": {"A": "bought", "B": "borrowed", "C": "sold", "D": "lost"},
                "correct_answer": "A",
                "correct_word": "bought"
            }
        ]
    }
}
```

如果无法找到C篇或D篇，passages数组中可以只包含找到的文章。如果完全找不到，passages为空数组。
如果找不到某篇文章的题目，questions数组可以为空。
如果找不到完形填空，cloze.found设为false，其他字段可省略。"""

    EXTRACT_WRITING_PROMPT = """你是一个英语试卷解析专家。请分析这份英语试卷文档，提取作文题目信息。

## 需要提取的信息

### 作文题目
- found: 是否找到作文题目（true/false）
- tasks: 作文小题数组。一份试卷的作文大题下面可能有两道或更多小题（例如“题目① / 题目② / 从下面两个题目中任选一题”），必须逐题拆开输出，不能把多道小题合并成一条。
- 每个 task 需要包含：
  - content: 该作文小题的完整内容（包括该小题自己的题干、提示词、提示问题、已给开头等）
  - requirements: 该小题适用的具体写作要求（如字数、格式等；如果是整道作文大题共享要求，也要复制到每个小题）
  - word_limit: 字数限制（如"80-100词"、"不少于60词"等）
  - writing_type: 文体类型
  - application_type: 应用文子类型（如果适用）
  - task_label: 小题标签（如"题目①"、"题目②"、"第一题"；如果没有则留空）

### 文体类型识别
- 每个 task 的 writing_type 必须从以下三个选项中选择一个：
  - "应用文"：书信、通知、邀请、邮件、回复、活动介绍、请假条、感谢信等
  - "记叙文"：讲述经历、描述事件、讲故事等
  - "其他"：无法归类或以上两者都不符合的

- 每个 task 的 application_type 仅当 writing_type 为"应用文"时填写，具体子类型如：
  - 书信、通知、邀请函、电子邮件、回复信、活动介绍、日记、便条等
  - 如果不是应用文，此字段留空

## 注意事项
1. 仔细识别作文题目的完整内容，包括所有背景信息和要求
2. 文体类型必须准确分类，根据题目要求判断
3. 如果作文大题下有两道或以上小题，必须在 tasks 中逐条返回，不能只返回其中一道
4. 如果找不到作文题目，将 found 设为 false

## 输出格式

请严格按照以下JSON格式输出，不要添加任何其他文字：

```json
{
    "found": true,
    "tasks": [
        {
            "task_label": "题目①",
            "content": "假设你是李华，你的英国笔友Tom对中国传统文化很感兴趣...",
            "requirements": "1. 词数80-100；2. 可适当增加细节...",
            "word_limit": "80-100词",
            "writing_type": "应用文",
            "application_type": "书信"
        },
        {
            "task_label": "题目②",
            "content": "某英文网站正在开展以“我梦想中的学校”为主题的征文活动...",
            "requirements": "1. 词数80-100；2. 可适当增加细节...",
            "word_limit": "80-100词",
            "writing_type": "记叙文",
            "application_type": ""
        }
    ]
}
```

如果未找到作文题目：
```json
{
    "found": false
}
```"""

    def __init__(self):
        self.api_key = settings.DASHSCOPE_API_KEY
        if not self.api_key:
            raise ValueError("DASHSCOPE_API_KEY未配置")

    def _build_file_messages(self, fileid: str, prompt: str) -> List[Dict[str, str]]:
        """构造 fileid 解析消息。"""
        return [
            {"role": "system", "content": "You are a helpful assistant specialized in parsing English exam papers."},
            {"role": "system", "content": f"fileid://{fileid}"},
            {"role": "user", "content": prompt},
        ]

    def _extract_docx_text(self, file_path: str) -> str:
        """提取 docx 文本，兼容段落和表格内容。"""
        doc = Document(file_path)

        parts: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    def _build_text_messages(self, file_path: str, prompt: str, max_chars: int = 30000) -> List[Dict[str, str]]:
        """构造文本回退解析消息。"""
        full_text = self._extract_docx_text(file_path)
        return [
            {"role": "system", "content": "You are a helpful assistant specialized in parsing English exam papers."},
            {"role": "user", "content": f"{prompt}\n\n## 文档内容\n\n{full_text[:max_chars]}"},
        ]

    def _should_fallback_to_text(self, error: Optional[str]) -> bool:
        """识别需要回退到文本解析的文件提取错误。"""
        if not error:
            return False

        lowered = error.lower()
        fallback_signals = [
            "encrypted or corrupted",
            "invalid_parameter_error",
            "file-extract",
            "unsupported",
            "failed to extract",
        ]
        return any(signal in lowered for signal in fallback_signals)

    async def _call_chat_completion(self, messages: List[Dict[str, str]], max_tokens: int) -> str:
        """统一调用 DashScope 聊天补全接口。"""
        result = await async_chat_completion(
            api_key=self.api_key,
            model="qwen-long",
            messages=messages,
            operation="llm_parser.chat_completion",
            temperature=0.1,
            max_tokens=max_tokens,
            timeout_seconds=120.0,
        )
        return result['choices'][0]['message']['content']

    async def upload_file(self, file_path: str) -> str:
        """
        上传文件到DashScope文件服务

        Args:
            file_path: 文件路径

        Returns:
            fileid: 文件ID
        """
        result = await async_upload_file(
            api_key=self.api_key,
            file_path=file_path,
            purpose="file-extract",
            operation="llm_parser.upload_file",
            timeout_seconds=120.0,
        )
        return result['id']

    async def parse_document(self, file_path: str, use_fileid: bool = True) -> LLMParseResult:
        """
        使用LLM解析文档

        Args:
            file_path: 文档路径
            use_fileid: 是否使用fileid方式（True）还是直接文本方式（False）

        Returns:
            LLMParseResult: 解析结果
        """
        try:
            if use_fileid:
                fileid = await self.upload_file(file_path)
                result = await self.parse_document_with_fileid(fileid, file_path=file_path)
                if result.success or not self._should_fallback_to_text(result.error):
                    return result
            else:
                pass

            messages = self._build_text_messages(file_path, self.EXTRACT_PROMPT)
            content = await self._call_chat_completion(messages, max_tokens=16000)
            return self._parse_llm_response(content)

        except Exception as e:
            return LLMParseResult(
                success=False,
                error=str(e)
            )

    def _parse_llm_response(self, content: str) -> LLMParseResult:
        """解析LLM返回的JSON内容"""
        try:
            json_str = self._extract_json_payload(content)
            data = self._load_json_payload(json_str)

            return LLMParseResult(
                success=True,
                metadata=data.get('metadata', {}),
                passages=data.get('passages', []),
                cloze=data.get('cloze', {}),
                raw_response=content
            )

        except json.JSONDecodeError as e:
            return LLMParseResult(
                success=False,
                error=f"JSON解析失败: {str(e)}",
                raw_response=content
            )

    def _extract_json_payload(self, content: str) -> str:
        """从模型返回内容中提取 JSON 主体。"""
        json_str = content.strip()

        if '```json' in content:
            start = content.find('```json') + 7
            end = content.find('```', start)
            json_str = content[start:end].strip()
        elif '```' in content:
            start = content.find('```') + 3
            end = content.find('```', start)
            json_str = content[start:end].strip()

        return json_str.lstrip("\ufeff")

    def _sanitize_json_string(self, json_str: str) -> str:
        """
        修复模型常见的脏 JSON：
        - 字符串中的裸换行/回车/制表符
        - 其他未转义控制字符
        """
        chars: List[str] = []
        in_string = False
        escaped = False

        for ch in json_str:
            if in_string:
                if escaped:
                    chars.append(ch)
                    escaped = False
                    continue

                if ch == "\\":
                    chars.append(ch)
                    escaped = True
                    continue

                if ch == '"':
                    chars.append(ch)
                    in_string = False
                    continue

                if ch == "\n":
                    chars.append("\\n")
                    continue
                if ch == "\r":
                    chars.append("\\r")
                    continue
                if ch == "\t":
                    chars.append("\\t")
                    continue
                if ord(ch) < 0x20:
                    chars.append(f"\\u{ord(ch):04x}")
                    continue

                chars.append(ch)
                continue

            chars.append(ch)
            if ch == '"':
                in_string = True

        return "".join(chars)

    def _load_json_payload(self, json_str: str) -> Dict:
        """优先直解，失败后做一次轻量 JSON 修复。"""
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as original_error:
            sanitized = self._sanitize_json_string(json_str)
            if sanitized != json_str:
                try:
                    return json.loads(sanitized)
                except json.JSONDecodeError:
                    pass
            raise original_error

    async def parse_document_with_fileid(self, fileid: str, file_path: Optional[str] = None) -> LLMParseResult:
        """
        使用已上传的fileid解析文档

        Args:
            fileid: 已上传到DashScope的文件ID
            file_path: 原始文件路径，fileid 解析失败时用于文本回退

        Returns:
            LLMParseResult: 解析结果
        """
        try:
            messages = self._build_file_messages(fileid, self.EXTRACT_PROMPT)
            content = await self._call_chat_completion(messages, max_tokens=16000)
            return self._parse_llm_response(content)

        except Exception as e:
            error = str(e)
            if file_path and self._should_fallback_to_text(error):
                try:
                    messages = self._build_text_messages(file_path, self.EXTRACT_PROMPT)
                    content = await self._call_chat_completion(messages, max_tokens=16000)
                    result = self._parse_llm_response(content)
                    if not result.success and result.error:
                        result.error = f"{result.error} (fileid回退到文本解析)"
                    return result
                except Exception as fallback_error:
                    error = f"{error} | 文本回退失败: {fallback_error}"

            return LLMParseResult(
                success=False,
                error=error
            )

    async def extract_writing(self, fileid: str, file_path: Optional[str] = None) -> "WritingExtractResult":
        """
        使用 LLM 提取作文题目信息

        Args:
            fileid: 已上传到 DashScope 的文件 ID
            file_path: 原始文件路径，fileid 解析失败时用于文本回退

        Returns:
            WritingExtractResult: 结构化提取结果
        """
        try:
            messages = self._build_file_messages(fileid, self.EXTRACT_WRITING_PROMPT)
            content = await self._call_chat_completion(messages, max_tokens=4000)
            result = self._parse_writing_response(content)
            if result.success and file_path:
                result = self._supplement_writing_tasks_from_docx(result, file_path)
            return result

        except Exception as e:
            error = str(e)
            if file_path and self._should_fallback_to_text(error):
                try:
                    messages = self._build_text_messages(file_path, self.EXTRACT_WRITING_PROMPT)
                    content = await self._call_chat_completion(messages, max_tokens=4000)
                    result = self._parse_writing_response(content)
                    if result.success:
                        result = self._supplement_writing_tasks_from_docx(result, file_path)
                    if not result.success and result.error:
                        result.error = f"{result.error} (fileid回退到文本解析)"
                    return result
                except Exception as fallback_error:
                    error = f"{error} | 文本回退失败: {fallback_error}"

            return WritingExtractResult(
                success=False,
                error=error
            )

    def _parse_writing_response(self, content: str) -> "WritingExtractResult":
        """解析作文提取响应"""
        try:
            json_str = self._extract_json_payload(content)
            data = self._load_json_payload(json_str)

            # 验证是否找到作文
            if not data.get('found', True):
                return WritingExtractResult(
                    success=False,
                    error="未找到作文题目"
                )

            raw_tasks = data.get("tasks")
            normalized_tasks: List[WritingExtractTask] = []

            if isinstance(raw_tasks, list) and raw_tasks:
                for item in raw_tasks:
                    if not isinstance(item, dict):
                        continue
                    normalized_tasks.extend(self._split_combined_writing_task(self._normalize_writing_task(item)))
            else:
                normalized_tasks.extend(self._split_combined_writing_task(self._normalize_writing_task(data)))

            deduped_tasks = dedupe_writing_tasks(normalized_tasks)

            if not deduped_tasks:
                return WritingExtractResult(
                    success=False,
                    error="未找到可用的作文小题"
                )

            return WritingExtractResult(
                success=True,
                tasks=deduped_tasks,
                raw_response=content,
            )

        except json.JSONDecodeError as e:
            return WritingExtractResult(
                success=False,
                error=f"JSON解析失败: {str(e)}"
            )

    def _normalize_writing_task(self, data: Dict) -> WritingExtractTask:
        """规范化单道作文任务。"""
        writing_type = data.get('writing_type', '其他')
        if writing_type not in ('应用文', '记叙文', '其他'):
            writing_type = '其他'

        application_type = data.get('application_type')
        if writing_type != '应用文':
            application_type = None
        elif application_type in ('null', ''):
            application_type = None

        return WritingExtractTask(
            content=(data.get('content') or '').strip(),
            requirements=(data.get('requirements') or '').strip(),
            word_limit=(data.get('word_limit') or '').strip(),
            writing_type=writing_type,
            application_type=application_type or "",
            task_label=(data.get('task_label') or '').strip(),
        )

    def _split_combined_writing_task(self, task: WritingExtractTask) -> List[WritingExtractTask]:
        """
        兼容模型把“题目① / 题目②”合并成一条的情况，按明显的小题标签拆分。
        """
        content = (task.content or "").strip()
        if not content:
            return []

        marker_pattern = re.compile(
            r'(?im)^\s*(题目\s*[①②③④⑤⑥⑦⑧⑨⑩]|题目\s*[1-9]|作文\s*[①②③④⑤⑥⑦⑧⑨⑩]|作文\s*[1-9]|第一题|第二题|第三题|Task\s*[1-9]|写作[一二三123])[\s：:]*'
        )
        matches = list(marker_pattern.finditer(content))
        if len(matches) < 2:
            return [task]

        shared_prefix = content[:matches[0].start()].strip()
        split_tasks: List[WritingExtractTask] = []

        for index, match in enumerate(matches):
            start = match.start()
            end = matches[index + 1].start() if index + 1 < len(matches) else len(content)
            segment = content[start:end].strip()
            if not segment:
                continue

            merged_content = f"{shared_prefix}\n{segment}".strip() if shared_prefix else segment
            split_tasks.append(
                WritingExtractTask(
                    content=merged_content,
                    requirements=task.requirements,
                    word_limit=task.word_limit,
                    writing_type=task.writing_type,
                    application_type=task.application_type,
                    task_label=match.group(1).strip(),
                )
            )

        return split_tasks or [task]

    def _supplement_writing_tasks_from_docx(
        self,
        result: WritingExtractResult,
        file_path: str,
    ) -> WritingExtractResult:
        """
        当 LLM 只返回一道作文，但原始文档里明显存在“题目① / 题目②”时，
        使用 DOCX 文本做兜底拆分，避免漏掉后续小题。
        """
        if not result.tasks or len(result.tasks) > 1:
            return result

        fallback_tasks = self._extract_writing_tasks_from_docx(file_path)
        if len(fallback_tasks) <= len(result.tasks):
            return result

        base_task = result.tasks[0]
        normalized_base = re.sub(r"\s+", " ", base_task.content or "").strip()
        supplemented: List[WritingExtractTask] = []

        for fallback_task in fallback_tasks:
            normalized_fallback = re.sub(r"\s+", " ", fallback_task.content or "").strip()
            if normalized_base and (
                normalized_base in normalized_fallback or normalized_fallback in normalized_base
            ):
                fallback_task.requirements = base_task.requirements or fallback_task.requirements
                fallback_task.word_limit = base_task.word_limit or fallback_task.word_limit
                fallback_task.writing_type = base_task.writing_type or fallback_task.writing_type
                fallback_task.application_type = base_task.application_type or fallback_task.application_type
            supplemented.append(fallback_task)

        return WritingExtractResult(
            success=True,
            tasks=dedupe_writing_tasks(supplemented),
            raw_response=result.raw_response,
        )

    def _extract_writing_tasks_from_docx(self, file_path: str) -> List[WritingExtractTask]:
        """从原始 docx 文本中兜底拆分作文小题。"""
        try:
            full_text = self._extract_docx_text(file_path)
        except Exception:
            return []

        if not full_text:
            return []

        section_text = self._extract_writing_section_from_text(full_text)
        if not section_text:
            return []

        matches = self._find_writing_task_markers(section_text)
        if len(matches) < 2:
            single_task = self._build_single_writing_task_from_section(section_text)
            return [single_task] if single_task else []

        shared_prefix = section_text[:matches[0].start()].strip()
        shared_requirements = shared_prefix
        shared_word_limit = self._extract_word_limit(shared_prefix)
        tasks: List[WritingExtractTask] = []
        seen_segments = set()
        seen_labels = set()

        for index, match in enumerate(matches):
            start = match.start()
            end = matches[index + 1].start() if index + 1 < len(matches) else len(section_text)
            segment = section_text[start:end].strip()
            if not segment or not self._looks_like_writing_prompt_segment(segment):
                continue

            merged_content = f"{shared_prefix}\n{segment}".strip() if shared_prefix else segment
            normalized_content = re.sub(r"\s+", " ", merged_content)
            if normalized_content in seen_segments:
                continue
            label = match.group(1).strip()
            if label in seen_labels:
                continue
            seen_segments.add(normalized_content)
            seen_labels.add(label)
            writing_type, application_type = self._infer_writing_type_from_text(segment)
            tasks.append(
                WritingExtractTask(
                    content=merged_content,
                    requirements=shared_requirements,
                    word_limit=shared_word_limit,
                    writing_type=writing_type,
                    application_type=application_type,
                    task_label=label,
                )
            )

        return dedupe_writing_tasks(tasks)

    def _build_single_writing_task_from_section(self, section_text: str) -> Optional[WritingExtractTask]:
        """当作文区只包含一道题时，直接按整段构造单题结果。"""
        normalized = re.sub(r"\s+", " ", section_text or "").strip()
        if not normalized or not self._looks_like_writing_prompt_segment(normalized[:600]):
            return None

        writing_type, application_type = self._infer_writing_type_from_text(normalized)
        return WritingExtractTask(
            content=section_text.strip(),
            requirements=section_text.strip(),
            word_limit=self._extract_word_limit(section_text),
            writing_type=writing_type,
            application_type=application_type,
            task_label="",
        )

    def _find_writing_task_markers(self, section_text: str) -> List[re.Match]:
        """查找作文小题起始标记，兼容题目①/②和老卷中的 15．/16． 这类编号格式。"""
        marker_pattern = re.compile(
            r'(?im)^\s*(题目\s*[①②③④⑤⑥⑦⑧⑨⑩]|题目\s*[1-9]|题目\s+[1-9]|作文\s*[①②③④⑤⑥⑦⑧⑨⑩]|作文\s*[1-9]|作文\s+[1-9]|第一题|第二题|第三题|Task\s*[1-9]|写作[一二三123])[\s：:]*'
        )
        matches = list(marker_pattern.finditer(section_text))
        if len(matches) >= 2:
            return matches

        # 兜底仅接受两位题号（如 19. / 39. / 52.），避免把题干内部的 1、2、3 小点误判成独立作文题。
        numbered_pattern = re.compile(r'(?im)^\s*(\d{2})[．\.、]\s*(?:（\d+分）)?')
        numbered_matches: List[re.Match] = []
        for match in numbered_pattern.finditer(section_text):
            line_end = section_text.find("\n", match.start())
            if line_end == -1:
                line_end = min(len(section_text), match.start() + 120)
            line_text = section_text[match.start():line_end].strip()
            if len(re.findall(r"[\u4e00-\u9fff]", line_text)) < 8 and "题目" not in line_text:
                continue
            if any(cue in line_text for cue in ("提示问题", "提示词", "一档文", "二档文", "三档文", "评分", "【解答】")):
                continue
            segment_end = section_text.find("\n", line_end + 1)
            if segment_end == -1:
                segment_end = min(len(section_text), match.start() + 220)
            preview = section_text[match.start():segment_end]
            if not self._looks_like_writing_prompt_segment(preview):
                continue
            numbered_matches.append(match)

        return numbered_matches

    def _looks_like_writing_prompt_segment(self, text: str) -> bool:
        normalized = re.sub(r"\s+", " ", text or "")
        if not normalized:
            return False

        prompt_cues = (
            "假设你是", "请用英语", "提示词语", "提示问题", "征文活动", "根据中文和英文提示",
            "根据所给提示", "写一封", "写一篇", "投稿", "活动举办的时间", "描述你梦想中的学校",
        )
        answer_cues = (
            "参考答案", "范文", "一档文", "二档文", "三档文", "评分标准", "要点齐全",
            "语言通顺", "扣分", "故选", "答案示例", "参考例文", "Yours,", "Best wishes", "Dear ", "Li Hua"
        )

        has_prompt_cue = any(cue in normalized for cue in prompt_cues)
        has_answer_cue = any(cue in normalized for cue in answer_cues)

        chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", normalized))
        english_words = len(re.findall(r"[A-Za-z]+(?:'[A-Za-z]+)?", normalized))

        if has_answer_cue:
            scoring_cues = ("一档文", "二档文", "三档文", "评分标准", "要点齐全", "语言通顺", "扣分", "答案示例")
            if any(cue in normalized for cue in scoring_cues):
                return False
        if re.search(r"[\(（]\d+[\)）]\s*[A-Da-d]\b", normalized):
            return False

        if has_prompt_cue:
            return True
        if has_answer_cue and chinese_chars < 20 and english_words > 40:
            return False
        return chinese_chars >= 20

    def _extract_writing_section_from_text(self, full_text: str) -> str:
        """从整份文档文本中截取作文区域。"""
        if not full_text:
            return ""

        section_cutoff_patterns = [
            r'(?im)^\s*(?:单项选择|单项填空|选择填空|完形填空|阅读理解|阅读短文|任务型阅读|语法填空|词汇运用|根据短文内容回答问题|补全对话|听力理解)\b',
            r'(?im)^\s*[一二三四五六七八九十百]+[、\.．]\s*(?:单项选择|选择填空|完形填空|阅读理解|阅读短文|任务型阅读|语法填空|词汇运用|根据短文内容回答问题|补全对话|听力理解)\b',
            r'(?im)^\s*\d+\s*[、\.．]\s*(?:单项选择|选择填空|完形填空|阅读理解|阅读短文|任务型阅读|语法填空|词汇运用)\b',
            r'(?im)^\s*(?:十一|十二|十三|十四|十五)[、\.．]\s*',
        ]

        patterns = [
            r'(?is)(?:五[、\.．]\s*文段表达|文段表达|书面表达|作文|写作)(.*)$',
        ]
        for pattern in patterns:
            match = re.search(pattern, full_text)
            if match:
                prefix = match.group(0).strip()
                duplicate_heading = re.search(
                    r'(?is)\n\s*(?:[一二三四五六七八九十]+[、\.．]\s*)?(?:书面表达|文段表达)',
                    prefix[200:],
                )
                if duplicate_heading:
                    prefix = prefix[:200 + duplicate_heading.start()].strip()
                answer_cutoff = re.search(
                    r'(?is)\n\s*(?:参考答案|答案及评分|评分标准|试题解析)\b',
                    prefix,
                )
                if answer_cutoff:
                    prefix = prefix[:answer_cutoff.start()].strip()
                for cutoff_pattern in section_cutoff_patterns:
                    section_cutoff = re.search(cutoff_pattern, prefix[80:])
                    if section_cutoff:
                        prefix = prefix[:80 + section_cutoff.start()].strip()
                        break
                return prefix
        return full_text

    def _extract_word_limit(self, text: str) -> str:
        if not text:
            return ""
        match = re.search(
            r'((?:不少于|不低于|约|至少)?\s*\d+\s*(?:[-~—至到]\s*\d+)?\s*词)',
            text,
        )
        return match.group(1).replace(" ", "") if match else ""

    def _infer_writing_type_from_text(self, text: str) -> tuple[str, str]:
        normalized = re.sub(r"\s+", " ", text or "")
        if re.search(r'(写一封|邮件|回信|书信|通知|倡议书|演讲稿|发言稿|邀请)', normalized):
            application_type = ""
            if "邀请" in normalized:
                application_type = "邀请"
            elif "回信" in normalized or "回复" in normalized:
                application_type = "回复信"
            elif "邮件" in normalized:
                application_type = "电子邮件"
            elif "通知" in normalized:
                application_type = "通知"
            elif "演讲稿" in normalized or "发言稿" in normalized:
                application_type = "演讲稿"
            elif "书信" in normalized or "信" in normalized:
                application_type = "书信"
            return "应用文", application_type

        if re.search(r'(征文|投稿|短文|经历|故事|梦想中的学校|描述|讲述)', normalized):
            return "记叙文", ""

        return "其他", ""


# 使用示例
async def test_llm_parser():
    """测试LLM解析器"""
    parser = LLMDocumentParser()

    test_file = "/Users/maoyu/Downloads/personal/项目-code/teachtools/2022-2025北京市各区各学校试题汇总/2022/初二春季期末/朝阳/2022北京朝阳初二（下）期末英语（教师版）.docx"

    print("正在上传文件并解析...")
    result = await parser.parse_document(test_file, use_fileid=True)

    if result.success:
        print("=== 元数据 ===")
        print(json.dumps(result.metadata, ensure_ascii=False, indent=2))

        print("\n=== 文章 ===")
        for p in result.passages:
            print(f"\n{p['passage_type']}篇 ({p['word_count']}词):")
            print(p['content'][:200] + "...")
    else:
        print(f"解析失败: {result.error}")
        if result.raw_response:
            print(f"原始响应: {result.raw_response[:500]}")


if __name__ == "__main__":
    import asyncio
    asyncio.run(test_llm_parser())
