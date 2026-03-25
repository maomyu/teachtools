"""
讲义 DOCX 导出服务

[INPUT]: 依赖 python-docx 和讲义聚合后的结构化数据
[OUTPUT]: 对外提供阅读/完形/作文讲义的 Word 导出能力
[POS]: backend/app/services 的讲义导出服务
[PROTOCOL]: 变更时更新此头部，然后检查 CLAUDE.md
"""
from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any, Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


class HandoutDocxExporter:
    """将讲义数据导出为可编辑的 DOCX。"""

    def build_reading_grade_docx(
        self,
        handout: dict[str, Any],
        paper_ids: Optional[list[int]] = None
    ) -> bytes:
        grade = handout.get("grade", "")
        edition = handout.get("edition", "teacher")
        document = self._create_document()

        self._add_cover(
            document,
            title=f"{grade}阅读CD篇讲义",
            edition=edition,
            paper_ids=paper_ids,
        )
        self._add_topics_overview(document, handout.get("topics", []), count_label="篇文章")

        for index, topic_content in enumerate(handout.get("content", []), start=1):
            self._add_topic_title(document, index, topic_content.get("topic", "未命名话题"))
            self._add_article_sources(document, topic_content.get("part1_article_sources", []))
            self._add_vocabulary_table(document, topic_content.get("part2_vocabulary", []))
            self._add_reading_passages(
                document,
                topic_content.get("part3_passages", []),
                edition,
            )

        return self._to_bytes(document)

    def build_cloze_grade_docx(
        self,
        handout: dict[str, Any],
        paper_ids: Optional[list[int]] = None
    ) -> bytes:
        grade = handout.get("grade", "")
        edition = handout.get("edition", "teacher")
        document = self._create_document()

        self._add_cover(
            document,
            title=f"{grade}完形填空讲义",
            edition=edition,
            paper_ids=paper_ids,
        )
        self._add_topics_overview(document, handout.get("topics", []), count_label="篇文章")

        for index, topic_content in enumerate(handout.get("content", []), start=1):
            self._add_topic_title(document, index, topic_content.get("topic", "未命名话题"))
            self._add_article_sources(document, topic_content.get("part1_article_sources", []))
            self._add_vocabulary_table(document, topic_content.get("part2_vocabulary", []))
            self._add_cloze_points_summary(
                document,
                topic_content.get("part3_points_by_type", {}),
                edition,
            )
            self._add_cloze_passages(
                document,
                topic_content.get("part4_passages", []),
                edition,
            )

        return self._to_bytes(document)

    def build_writing_grade_docx(
        self,
        handout: dict[str, Any],
        paper_ids: Optional[list[int]] = None
    ) -> bytes:
        grade = handout.get("grade", "")
        edition = handout.get("edition", "teacher")
        document = self._create_document()

        self._add_cover(
            document,
            title=f"{grade}作文讲义",
            edition=edition,
            paper_ids=paper_ids,
        )
        self._add_topics_overview(handout=document, topics=handout.get("topics", []), count_label="道题")

        for index, topic_content in enumerate(handout.get("content", []), start=1):
            self._add_topic_title(document, index, topic_content.get("topic", "未命名话题"))
            self._add_writing_topic_stats(document, topic_content.get("part1_topic_stats", {}))
            self._add_writing_frameworks(document, topic_content.get("part2_frameworks", []))
            self._add_writing_expressions(document, topic_content.get("part3_expressions", []))
            self._add_writing_samples(
                document,
                topic_content.get("part4_samples", []),
                edition,
            )

        return self._to_bytes(document)

    @staticmethod
    def build_filename(
        base_title: str,
        edition: str,
        paper_ids: Optional[list[int]] = None
    ) -> str:
        scope_suffix = f"_选{len(paper_ids)}卷" if paper_ids else ""
        edition_label = "教师版" if edition == "teacher" else "学生版"
        date_suffix = datetime.now().strftime("%Y-%m-%d")
        return f"{base_title}{scope_suffix}_{edition_label}_{date_suffix}.docx"

    def _create_document(self) -> Document:
        document = Document()
        section = document.sections[0]
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(10.5)
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        for style_name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
            style = document.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        return document

    def _to_bytes(self, document: Document) -> bytes:
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _add_cover(
        self,
        document: Document,
        title: str,
        edition: str,
        paper_ids: Optional[list[int]]
    ) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run("教师版" if edition == "teacher" else "学生版").font.size = Pt(14)

        scope_text = f"已选 {len(paper_ids)} 份试卷" if paper_ids else "范围：全年级全部试卷"
        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(scope_text)
        meta.add_run(f"\n生成日期：{datetime.now().strftime('%Y-%m-%d')}")

        document.add_page_break()

    def _add_topics_overview(
        self,
        handout: Document,
        topics: list[dict[str, Any]],
        count_label: str
    ) -> None:
        handout.add_heading("目录", level=1)
        if not topics:
            handout.add_paragraph("暂无讲义内容。")
            handout.add_page_break()
            return

        for index, topic in enumerate(topics, start=1):
            recent_years = topic.get("recent_years") or []
            years_text = f" | 近年：{', '.join(str(year) for year in recent_years)}" if recent_years else ""
            handout.add_paragraph(
                f"{index}. {topic.get('topic', '未命名话题')} | {topic.get('passage_count', topic.get('task_count', 0))}{count_label}{years_text}",
                style="List Bullet",
            )

        handout.add_page_break()

    def _add_topic_title(self, document: Document, index: int, topic: str) -> None:
        document.add_heading(f"{index}. {topic}", level=1)

    def _add_section_heading(self, document: Document, title: str) -> None:
        document.add_heading(title, level=2)

    def _add_article_sources(self, document: Document, sources: list[dict[str, Any]]) -> None:
        self._add_section_heading(document, "Part 1 文章来源")
        if not sources:
            document.add_paragraph("暂无文章来源。")
            return

        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["年份", "地区", "考试", "学期", "篇目"]
        for cell, title in zip(table.rows[0].cells, headers):
            cell.text = title

        for source in sources:
            row = table.add_row().cells
            row[0].text = self._safe_text(source.get("year"))
            row[1].text = self._safe_text(source.get("region"))
            row[2].text = self._safe_text(source.get("exam_type"))
            row[3].text = self._safe_text(source.get("semester"))
            row[4].text = "；".join(
                self._format_passage_ref(passage) for passage in source.get("passages", [])
            )

        document.add_paragraph()

    def _add_vocabulary_table(self, document: Document, vocabulary: list[dict[str, Any]]) -> None:
        self._add_section_heading(document, "Part 2 高频词汇")
        if not vocabulary:
            document.add_paragraph("暂无高频词汇。")
            return

        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["单词", "音标", "释义", "频次", "来源"]
        for cell, title in zip(table.rows[0].cells, headers):
            cell.text = title

        for item in vocabulary:
            row = table.add_row().cells
            row[0].text = self._safe_text(item.get("word"))
            row[1].text = self._safe_text(item.get("phonetic"))
            row[2].text = self._safe_text(item.get("definition"))
            row[3].text = self._safe_text(item.get("frequency"))
            row[4].text = self._map_source_type(item.get("source_type"))

        document.add_paragraph()

    def _add_reading_passages(
        self,
        document: Document,
        passages: list[dict[str, Any]],
        edition: str
    ) -> None:
        self._add_section_heading(document, "Part 3 阅读文章")
        if not passages:
            document.add_paragraph("暂无阅读文章。")
            return

        for index, passage in enumerate(passages, start=1):
            document.add_heading(
                f"文章 {index} · {self._safe_text(passage.get('type'))}篇 {self._safe_text(passage.get('title'))}".strip(),
                level=3,
            )
            document.add_paragraph(self._format_source_line(passage.get("source")))
            self._add_multiline_text(document, passage.get("content"))

            questions = passage.get("questions", [])
            if questions:
                document.add_paragraph("题目：")
                for question in questions:
                    q_text = f"{self._safe_text(question.get('number'))}. {self._safe_text(question.get('text'))}".strip()
                    document.add_paragraph(q_text, style="List Number")
                    for option_key in ("A", "B", "C", "D"):
                        option_value = (question.get("options") or {}).get(option_key)
                        if option_value:
                            document.add_paragraph(f"{option_key}. {option_value}", style="List Bullet")
                    if edition == "teacher":
                        answer = self._safe_text(question.get("correct_answer"))
                        explanation = self._safe_text(question.get("explanation"))
                        if answer:
                            document.add_paragraph(f"答案：{answer}")
                        if explanation:
                            document.add_paragraph(f"解析：{explanation}")

            if index < len(passages):
                document.add_page_break()

    def _add_cloze_points_summary(
        self,
        document: Document,
        points_by_type: dict[str, Any],
        edition: str
    ) -> None:
        self._add_section_heading(document, "Part 3 考点分布")
        if not points_by_type:
            document.add_paragraph("暂无考点分布。")
            return

        for code in sorted(points_by_type.keys()):
            group = points_by_type.get(code) or {}
            document.add_heading(
                f"{code} {self._safe_text(group.get('name'))} | {self._safe_text(group.get('category_name'))}",
                level=3,
            )
            for point in group.get("points", []):
                summary = f"{self._safe_text(self._lookup(point, 'word'))} · 出现 {self._safe_text(self._lookup(point, 'frequency'))} 次"
                if self._lookup(point, "phrase"):
                    summary += f" · 搭配：{self._safe_text(self._lookup(point, 'phrase'))}"
                elif self._lookup(point, "definition"):
                    summary += f" · 释义：{self._safe_text(self._lookup(point, 'definition'))}"
                document.add_paragraph(summary, style="List Bullet")

                if edition == "teacher":
                    if self._lookup(point, "textbook_meaning") or self._lookup(point, "context_meaning"):
                        document.add_paragraph(
                            f"课本义：{self._safe_text(self._lookup(point, 'textbook_meaning'))} | 语境义：{self._safe_text(self._lookup(point, 'context_meaning'))}"
                        )
                    occurrences = self._lookup(point, "occurrences") or []
                    if occurrences:
                        first = occurrences[0]
                        document.add_paragraph(f"示例：{self._safe_text(self._lookup(first, 'sentence'))}")
                        analysis = self._lookup(first, "analysis")
                        if self._lookup(analysis, "explanation"):
                            document.add_paragraph(f"解析：{self._safe_text(self._lookup(analysis, 'explanation'))}")

            document.add_paragraph()

    def _add_cloze_passages(
        self,
        document: Document,
        passages: list[dict[str, Any]],
        edition: str
    ) -> None:
        self._add_section_heading(document, "Part 4 完形原文")
        if not passages:
            document.add_paragraph("暂无完形文章。")
            return

        for index, passage in enumerate(passages, start=1):
            document.add_heading(f"完形 {index}", level=3)
            document.add_paragraph(self._format_source_line(self._lookup(passage, "source")))
            self._add_multiline_text(document, self._lookup(passage, "content"))

            if edition == "teacher":
                document.add_paragraph("空格考点：")
                for point in sorted(self._lookup(passage, "points") or [], key=lambda item: self._lookup(item, "blank_number") or 0):
                    summary = (
                        f"第 {self._safe_text(self._lookup(point, 'blank_number'))} 空 | "
                        f"答案：{self._safe_text(self._lookup(point, 'correct_answer'))} / {self._safe_text(self._lookup(point, 'correct_word'))} | "
                        f"考点：{self._safe_text(self._lookup(point, 'primary_point_code') or self._lookup(point, 'point_type'))}"
                    )
                    document.add_paragraph(summary, style="List Bullet")
                    if self._lookup(point, "translation"):
                        document.add_paragraph(f"释义：{self._safe_text(self._lookup(point, 'translation'))}")
                    if self._lookup(point, "explanation"):
                        document.add_paragraph(f"解析：{self._safe_text(self._lookup(point, 'explanation'))}")
                    rejection_points = self._lookup(point, "rejection_points") or []
                    if rejection_points:
                        reasons = "；".join(
                            f"{self._safe_text(self._lookup(item, 'option_word'))}: {self._safe_text(self._lookup(item, 'rejection_reason') or self._lookup(item, 'explanation'))}"
                            for item in rejection_points
                        )
                        document.add_paragraph(f"排错：{reasons}")

            if index < len(passages):
                document.add_page_break()

    def _add_writing_topic_stats(self, document: Document, stats: dict[str, Any]) -> None:
        self._add_section_heading(document, "Part 1 话题概览")
        document.add_paragraph(
            f"题目数：{self._safe_text(stats.get('task_count'))} | 范文数：{self._safe_text(stats.get('sample_count'))} | 近年：{', '.join(str(year) for year in (stats.get('recent_years') or []))}"
        )
        document.add_paragraph()

    def _add_writing_frameworks(self, document: Document, frameworks: list[dict[str, Any]]) -> None:
        self._add_section_heading(document, "Part 2 写作框架")
        if not frameworks:
            document.add_paragraph("暂无写作框架。")
            return

        for framework in frameworks:
            document.add_heading(self._safe_text(framework.get("writing_type")), level=3)
            for section in framework.get("sections", []):
                document.add_paragraph(
                    f"{self._safe_text(section.get('name'))}：{self._safe_text(section.get('description'))}",
                    style="List Bullet",
                )
                for example in section.get("examples", []) or []:
                    document.add_paragraph(f"示例：{example}")

        document.add_paragraph()

    def _add_writing_expressions(self, document: Document, expressions: list[dict[str, Any]]) -> None:
        self._add_section_heading(document, "Part 3 高频表达")
        if not expressions:
            document.add_paragraph("暂无高频表达。")
            return

        for group in expressions:
            document.add_heading(self._safe_text(group.get("category")), level=3)
            for item in group.get("items", []) or []:
                document.add_paragraph(str(item), style="List Bullet")

        document.add_paragraph()

    def _add_writing_samples(
        self,
        document: Document,
        samples: list[dict[str, Any]],
        edition: str
    ) -> None:
        self._add_section_heading(document, "Part 4 范文展示")
        if not samples:
            document.add_paragraph("暂无范文。")
            return

        for index, sample in enumerate(samples, start=1):
            document.add_heading(f"范文 {index}", level=3)
            document.add_paragraph(f"题目：{self._safe_text(sample.get('task_content'))}")
            source = sample.get("source") or {}
            source_parts = [
                self._safe_text(source.get("year")),
                self._safe_text(source.get("region")),
                self._safe_text(source.get("exam_type")),
                self._safe_text(source.get("semester")),
            ]
            source_line = " · ".join(part for part in source_parts if part)
            meta_parts = [part for part in [source_line, f"{self._safe_text(sample.get('word_count'))} 词"] if part.strip()]
            if meta_parts:
                document.add_paragraph("来源：" + " | ".join(meta_parts))
            self._add_multiline_text(document, sample.get("sample_content"))

            if edition == "teacher" and sample.get("translation"):
                document.add_paragraph("译文：")
                self._add_multiline_text(document, sample.get("translation"))

            highlights = sample.get("highlighted_sentences") or []
            if edition == "teacher" and highlights:
                document.add_paragraph("亮点句：")
                for item in highlights:
                    document.add_paragraph(
                        f"{self._safe_text(item.get('highlight_type'))} | {self._safe_text(item.get('sentence'))}",
                        style="List Bullet",
                    )
                    if item.get("explanation"):
                        document.add_paragraph(f"说明：{item['explanation']}")

            if index < len(samples):
                document.add_page_break()

    def _add_multiline_text(self, document: Document, text: Any) -> None:
        cleaned = self._safe_text(text)
        if not cleaned:
            return
        for line in cleaned.splitlines():
            line = line.strip()
            if line:
                document.add_paragraph(line)

    def _format_passage_ref(self, passage: dict[str, Any]) -> str:
        passage_type = self._safe_text(self._lookup(passage, "type"))
        title = self._safe_text(self._lookup(passage, "title"))
        passage_id = self._safe_text(self._lookup(passage, "id"))
        parts = [part for part in [passage_type, title, f"ID {passage_id}" if passage_id else ""] if part]
        return " ".join(parts) if parts else "未命名篇目"

    def _format_source_line(self, source: Optional[dict[str, Any]]) -> str:
        source = source or {}
        parts = [
            self._safe_text(self._lookup(source, "year")),
            self._safe_text(self._lookup(source, "region")),
            self._safe_text(self._lookup(source, "grade")),
            self._safe_text(self._lookup(source, "semester")),
            self._safe_text(self._lookup(source, "exam_type")),
            self._safe_text(self._lookup(source, "filename")),
        ]
        filtered = [part for part in parts if part]
        return "来源：" + " | ".join(filtered) if filtered else "来源：未标注"

    def _map_source_type(self, source_type: Any) -> str:
        mapping = {
            "both": "阅读+完形",
            "reading": "阅读",
            "cloze": "完形",
        }
        return mapping.get(str(source_type), self._safe_text(source_type))

    def _safe_text(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
        return str(value).strip()

    def _lookup(self, item: Any, key: str, default: Any = None) -> Any:
        if item is None:
            return default
        if isinstance(item, dict):
            return item.get(key, default)
        return getattr(item, key, default)
